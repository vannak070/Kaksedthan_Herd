import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re

def create_styled_docx(md_filepath, docx_filepath):
    doc = docx.Document()

    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Set Base Normal Style Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x24, 0x2A)

    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def set_table_borders(table):
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    def add_callout_box(doc, text, callout_type="note"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        
        bg_color = "F0F9FF" if callout_type == "note" else ("FFFBEB" if callout_type in ["warning", "caution"] else "F0FDF4")
        border_color = "0284C7" if callout_type == "note" else ("D97706" if callout_type in ["warning", "caution"] else "16A34A")

        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=120, bottom=120, left=200, right=200)

        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
            f'<w:top w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def render_markdown_table(doc, table_lines):
        rows_data = []
        for line in table_lines:
            if re.match(r'^\s*\|?\s*:?-+:?\s*\|', line):
                continue
            cols = [c.strip() for c in line.strip().strip('|').split('|')]
            if cols and any(cols):
                rows_data.append(cols)

        if not rows_data:
            return

        num_rows = len(rows_data)
        num_cols = max(len(r) for r in rows_data)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        for r_idx, row_cols in enumerate(rows_data):
            row = table.rows[r_idx]
            is_header = (r_idx == 0)
            for c_idx in range(num_cols):
                if c_idx < len(row_cols):
                    cell_text = row_cols[c_idx]
                    cell = row.cells[c_idx]
                    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

                    if is_header:
                        set_cell_background(cell, "0F172A")
                    elif r_idx % 2 == 1:
                        set_cell_background(cell, "F8FAFC")
                    else:
                        set_cell_background(cell, "FFFFFF")

                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)

                    # Parse bold & code in table cells
                    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', cell_text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            r = p.add_run(part[2:-2])
                            r.bold = True
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if is_header else RGBColor(0x0F, 0x17, 0x2A)
                        elif part.startswith('`') and part.endswith('`'):
                            r = p.add_run(part[1:-1])
                            r.font.name = 'Consolas'
                            r.font.size = Pt(9.5)
                            r.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8) if is_header else RGBColor(0x02, 0x84, 0xC7)
                        else:
                            r = p.add_run(part)
                            if is_header:
                                r.bold = True
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            else:
                                r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block handler
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "0F172A")
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                code_text = "".join(code_block_lines)
                r = p.add_run(code_text)
                r.font.name = 'Consolas'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                code_block_lines = []
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Markdown Table handler
        if '|' in line and not line.strip().startswith('>'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                render_markdown_table(doc, table_lines)
                table_lines = []

        line_str = line.strip()

        if not line_str:
            i += 1
            continue

        # Headings
        if line_str.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(line_str[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(22)
            run.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line_str.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line_str[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        elif line_str.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line_str[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        # Callouts / Alert boxes
        elif line_str.startswith('> [!IMPORTANT]') or line_str.startswith('> [!WARNING]') or line_str.startswith('> [!CAUTION]') or line_str.startswith('> [!NOTE]'):
            callout_type = "note"
            if "WARNING" in line_str or "CAUTION" in line_str:
                callout_type = "warning"
            callout_lines = []
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>'):
                callout_lines.append(lines[i].strip().lstrip('> ').strip())
                i += 1
            add_callout_box(doc, " ".join(callout_lines), callout_type=callout_type)
            continue

        # Bullet lists / Checklists
        elif line_str.startswith('* ') or line_str.startswith('- ') or line_str.startswith('- [ ] ') or line_str.startswith('- [x] '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)

            text_content = line_str
            if line_str.startswith('- [x] ') or line_str.startswith('- [ ] '):
                prefix = "☒ " if line_str.startswith('- [x] ') else "☐ "
                text_content = prefix + line_str[6:]
            else:
                text_content = line_str[2:]

            parts = re.split(r'(\*\*.*?\*\*|`.*?`|\$\$.*?\$\$|\$.*?\$)', text_content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Consolas'
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
                else:
                    p.add_run(part)

        # Standard Paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)

            parts = re.split(r'(\*\*.*?\*\*|`.*?`|\$\$.*?\$\$|\$.*?\$)', line_str)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Consolas'
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
                elif part.startswith('$') and part.endswith('$'):
                    r = p.add_run(part.strip('$'))
                    r.font.name = 'Cambria Math'
                    r.italic = True
                else:
                    p.add_run(part)

        i += 1

    if in_table:
        render_markdown_table(doc, table_lines)

    doc.save(docx_filepath)
    print(f"✅ Document successfully created at: {docx_filepath}")

if __name__ == '__main__':
    create_styled_docx(
        '/Users/vannakath/HOVA_Project/Kaksethan_Herdbook/OFFICIAL_DEVELOPMENT_AND_DEPLOYMENT_GUIDE.md',
        '/Users/vannakath/HOVA_Project/Kaksethan_Herdbook/OFFICIAL_DEVELOPMENT_AND_DEPLOYMENT_GUIDE.docx'
    )
